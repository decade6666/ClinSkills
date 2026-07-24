"""utils/output_docx.py — docx 三线表输出

从 output_format.py 按介质拆分而来（docx 部分）。公开函数 save_table_to_docx_threeline；
其余 set_*/calculate_column_widths/merge_cells_vertical 为其辅助。
既有代码经 output_format.py 聚合入口导入即可，无需改动。
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_border(cell, **kwargs):
    """
    设置单元格边框
    
    参数示例:
    top={"sz": 12, "val": "single", "color": "000000"}
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def _apply_row_borders(cell, col_idx, num_cols, top, bottom, dashed_border, no_border):
    """三线表单元格边框：首列左无框、末列右无框，其余侧边虚线；top/bottom 由调用方按行位给。

    与原逐列分支完全等效（单列时按「首列」处理：左无框、右虚线）。
    """
    left = no_border if col_idx == 0 else dashed_border
    right = no_border if (col_idx == num_cols - 1 and col_idx != 0) else dashed_border
    set_cell_border(cell, top=top, bottom=bottom, left=left, right=right)

def set_cell_background(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_run_font(run, chinese_font='宋体', english_font='Times New Roman', size=10.5, bold=False):
    """设置run的字体：中文字体，英文和数字字体"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = english_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
    run._element.rPr.rFonts.set(qn('w:ascii'), english_font)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), english_font)

def set_cell_font(cell, size=10.5, bold=False):
    """设置单元格字体：中文宋体，英文和数字Times New Roman"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold)

def set_row_height(row, height_cm=0.6):
    """设置行的最小高度"""
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # 转换为twips (1cm = 567 twips)
    trHeight.set(qn('w:hRule'), 'atLeast')  # 最小行高
    trPr.append(trHeight)

def set_table_width(table, width_cm):
    """设置表格总宽度"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # 设置表格宽度
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_cm * 567)))  # 转换为twips
    tblW.set(qn('w:type'), 'dxa')  # dxa表示使用绝对值
    tblPr.append(tblW)

def calculate_column_widths(df, available_width_cm):
    """
    根据内容自动计算列宽，并确保总宽度等于可用宽度
    
    参数：
    df : pandas.DataFrame
        数据框
    available_width_cm : float
        可用页面宽度（厘米）

    返回：
    list : 每列的宽度（厘米）
    """
    widths = []
    
    # 首先计算每列的相对权重
    for col in df.columns:
        # 计算列名长度
        max_len = len(str(col))
        
        # 计算列数据的最大长度
        for value in df[col]:
            max_len = max(max_len, len(str(value)))
        
        # 估算宽度：每字符约 0.35cm + 0.5cm 基准（未区分中英文）
        width = max_len * 0.35 + 0.5
        width = max(1.5, width)  # 最小宽度1.5cm
        
        widths.append(width)
    
    # 计算总权重
    total_weight = sum(widths)
    
    # 按比例分配可用宽度
    widths = [(w / total_weight) * available_width_cm for w in widths]
    
    return widths

def merge_cells_vertical(table, start_row, end_row, col_idx):
    """
    垂直合并单元格
    
    参数:
        table: Word表格对象
        start_row: 起始行索引
        end_row: 结束行索引
        col_idx: 列索引
    """
    # 获取起始单元格
    start_cell = table.rows[start_row].cells[col_idx]
    
    # 保存第一个单元格的文本
    original_text = start_cell.text
    
    # 清空其他单元格的文本（合并前）
    for row_idx in range(start_row + 1, end_row + 1):
        table.rows[row_idx].cells[col_idx].text = ""
    
    # 合并单元格
    for row_idx in range(start_row + 1, end_row + 1):
        end_cell = table.rows[row_idx].cells[col_idx]
        start_cell.merge(end_cell)
    
    # 恢复第一个单元格的文本（防止合并时文本被拼接）
    start_cell.text = original_text
    
    # 设置合并后单元格的垂直居中对齐
    tc = start_cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    
    # 重新设置字体格式（因为清空和重新赋值后格式可能丢失）
    for paragraph in start_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, size=10.5, bold=False)

def save_table_to_docx_threeline(df: pd.DataFrame, output_path: str, title: str, notes: list = None, 
                                  row_height_cm=0.6, auto_width=True, include_notes=True, merge_columns=None, alignment = WD_TABLE_ALIGNMENT.CENTER):
    """
    将 DataFrame 保存为标准三线表格式的 .docx 文件
    
    参数：
    df : pandas.DataFrame
        要保存的表格数据
    output_path : str
        输出 .docx 文件的路径
    title : str
        表格的标题
    notes : list, optional
        表格下方的注释列表，每项为一条注释
    row_height_cm : float, optional
        最小行高（厘米），默认0.6cm
    auto_width : bool, optional
        是否自动调整列宽，默认True
    include_notes : bool, optional
        是否输出脚注，默认True
    merge_columns : list, optional
        需要合并的列名列表，当上下单元格值相同时进行合并，默认None
    alignment : optional
        表格整体对齐方式，默认居中（WD_TABLE_ALIGNMENT.CENTER）
    """
    doc = Document()
    
    # 设置页面为横向A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    
    # 设置为A4纸张大小（横向）
    # A4纸张：宽21cm，高29.7cm
    # 横向时需要交换
    section.page_width = Cm(29.7)   # 横向时的宽度
    section.page_height = Cm(21)    # 横向时的高度
    
    # 设置页边距
    section.top_margin = Cm(2.5)     # 上边距 2.5cm
    section.bottom_margin = Cm(2)    # 下边距 2cm
    section.left_margin = Cm(2.5)    # 左边距 2.5cm
    section.right_margin = Cm(2)     # 右边距 2cm
    
    # 计算可用宽度（从 section 动态取值，避免与页边距硬编码漂移）
    available_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm  # = 25.2cm
    
    # 添加标题（居中，宋体，五号字）
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    set_run_font(title_run, chinese_font='宋体', english_font='Times New Roman', size=10.5, bold=True)
    
    # 设置标题后的间距
    title_para.paragraph_format.space_after = Pt(6)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # 设置表格居中对齐
    table.alignment = alignment
    
    # 设置表格总宽度
    set_table_width(table, available_width_cm)
    
    # 计算并设置列宽
    if auto_width:
        col_widths = calculate_column_widths(df, available_width_cm)
        # 对每一行的每一列设置宽度
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)
    
    # 边框配置
    # 0.5磅 = 4 sz (1磅 = 8 sz)
    thick_border = {"sz": 4, "val": "single", "color": "000000"}  # 0.5磅实线（顶线、底线、表头底线）
    # 使用dotted样式代替dashed，更密集
    dashed_border = {"sz": 4, "val": "dotted", "color": "000000"}  # 密集虚线（其他线）
    no_border = {"sz": 0, "val": "none"}
    
    # 设置表头行
    hdr_row = table.rows[0]
    set_row_height(hdr_row, row_height_cm)  # 设置表头行高
    
    hdr_cells = hdr_row.cells
    for i, column_name in enumerate(df.columns):
        cell = hdr_cells[i]
        cell.text = str(column_name)
        
        # 设置背景色（浅灰表头底色）
        set_cell_background(cell, 'AEAAAA')
        
        # 设置对齐
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 设置段落间距
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        
        # 设置字体（五号字=10.5磅）
        set_cell_font(cell, size=10.5, bold=True)
        
        # 设置边框：顶部0.5磅实线，底部0.5磅实线，左右虚线
        _apply_row_borders(cell, i, len(df.columns), thick_border, thick_border,
                           dashed_border, no_border)
    
    # 添加数据行
    for pos, (idx, row_data) in enumerate(df.iterrows()):
        data_row = table.add_row()
        set_row_height(data_row, row_height_cm)  # 设置数据行高
        
        # 设置新增行的列宽
        if auto_width:
            for i, width in enumerate(col_widths):
                data_row.cells[i].width = Cm(width)
        
        row_cells = data_row.cells
        is_last_row = (pos == len(df) - 1)
        
        for i, value in enumerate(row_data):
            cell = row_cells[i]
            cell.text = str(value)
            
            # 设置对齐
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 设置段落间距
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # 设置字体（五号字=10.5磅）
            set_cell_font(cell, size=10.5, bold=False)
            
            # 设置边框：末行底部0.5磅实线（三线表底线），其余虚线
            bottom = thick_border if is_last_row else dashed_border
            _apply_row_borders(cell, i, len(df.columns), dashed_border, bottom,
                               dashed_border, no_border)
    
    # 执行单元格合并（新增功能）
    if merge_columns:
        for col_name in merge_columns:
            if col_name not in df.columns:
                continue
            
            col_idx = df.columns.get_loc(col_name)
            
            # 找出需要合并的单元格范围
            merge_ranges = []
            start_row = 0  # 从第0行开始（DataFrame索引）
            current_value = df.iloc[0][col_name]
            
            for i in range(1, len(df)):
                if df.iloc[i][col_name] == current_value:
                    # 值相同，继续
                    continue
                else:
                    # 值不同，记录合并范围
                    if i > start_row + 1:  # 只有当范围大于1时才记录
                        merge_ranges.append((start_row, i - 1, col_idx))
                    start_row = i
                    current_value = df.iloc[i][col_name]
            
            # 处理最后一组
            if len(df) > start_row + 1:
                merge_ranges.append((start_row, len(df) - 1, col_idx))
            
            # 执行合并（Word表格行索引需要+1，因为第0行是表头）
            for start, end, col in merge_ranges:
                merge_cells_vertical(table, start + 1, end + 1, col)

                merged_cell = table.rows[start + 1].cells[col]
                for paragraph in merged_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加注释（小五号字=9磅）- 根据include_notes参数控制
    if include_notes and notes:
        notes_para = doc.add_paragraph()
        # 设置段落格式：段前段后均为0，单倍行距
        notes_para.paragraph_format.space_before = Pt(0)
        notes_para.paragraph_format.space_after = Pt(0)
        notes_para.paragraph_format.line_spacing = 1.0  # 单倍行距
        
        # 添加"注："并换行
        run = notes_para.add_run('注：\n')
        set_run_font(run, size=9, bold=False)  # 小五号字=9磅
        
        # 添加注释内容
        for i, note in enumerate(notes, 1):
            if i > 1:
                notes_para.add_run('\n')
            run = notes_para.add_run(f'{i}. {note}')
            set_run_font(run, size=9, bold=False)  # 小五号字=9磅
    
    # 保存文档（python-docx生成的文档默认就是可编辑的）
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    print(f"{title}已保存为 '{output_path}'")
